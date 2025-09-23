# app.py
import streamlit as st
import os, re, json, psutil
from datetime import datetime
from docx import Document
from docx.shared import Inches
from config import OUTPUT_DIR
from agent_core import ContentAgent
from local_llm_loader import load_local_llm
from image_generator import LocalImageGenerator
from data_visualizer import DataVisualizer

def save_document(filename, title, final_content, visual_mapping):
    if not os.path.exists(OUTPUT_DIR): os.makedirs(OUTPUT_DIR)
    doc = Document(); doc.add_heading(title, level=1)
    for section in final_content:
        doc.add_heading(section['topic'], level=2)
        # Use a more robust regex for splitting to handle content without figures
        parts = re.split(r'(<FIGURE>.*?</FIGURE>)', section.get('content', ''), flags=re.DOTALL)
        for part in parts:
            if not part.strip(): continue
            if part.startswith('<FIGURE>'):
                visual_match = re.search(r'(\[IMAGE\|.*?\]|\[CHART\|.*?\])', part, re.DOTALL)
                caption_match = re.search(r'<CAPTION>(.*?)</CAPTION>', part, re.DOTALL)
                if visual_match:
                    placeholder = visual_match.group(1)
                    if placeholder in visual_mapping and visual_mapping.get(placeholder):
                        try:
                            doc.add_picture(visual_mapping[placeholder], width=Inches(6.0))
                            if caption_match: doc.add_paragraph(caption_match.group(1).strip(), style='Caption')
                        except Exception as e: doc.add_paragraph(f"[Error adding visual: {e}]", style='Comment')
                    else: doc.add_paragraph(f"[Visual gen failed for: {placeholder}]", style='Comment')
            else:
                if part.strip().startswith('|') and '|' in part:
                    try:
                        lines = [l.strip() for l in part.strip().split('\n')]; header = [h.strip() for h in lines[0].strip('|').split('|')]
                        table = doc.add_table(rows=1, cols=len(header)); table.style = 'Table Grid'
                        for i, h in enumerate(header): table.rows[0].cells[i].text = h
                        for line in lines[2:]:
                            row_cells = table.add_row().cells
                            for i, cell in enumerate(line.strip('|').split('|')): row_cells[i].text = cell.strip()
                    except Exception as e: doc.add_paragraph(f"[Error rendering table: {e}]\n{part}")
                else: doc.add_paragraph(part)
    filepath = os.path.join(OUTPUT_DIR, f"{filename.replace(' ', '_').lower()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(filepath); return filepath

st.set_page_config(layout="wide"); st.title("🎓 On-Device Agentic Researcher")
for key in ['agent', 'image_agent', 'chart_agent', 'document_sections', 'outline', 'doc_title', 'metrics']:
    if key not in st.session_state: st.session_state[key] = None if key.endswith('_agent') else [] if key.endswith('_sections') else "" if 'title' in key or 'outline' in key else None

with st.sidebar:
    st.header("⚙️ Configuration")
    hardware_mode = st.radio("Select LLM Acceleration:", ("CPU", "Hybrid", "GPU"))
    if st.button("Load All On-Device Models"):
        try:
            with st.spinner("Loading Language Model..."):
                st.session_state.llm = load_local_llm(mode=hardware_mode.lower())
                st.session_state.agent = ContentAgent(st.session_state.llm)
            st.success("Language Model & Agent Loaded!")
            with st.spinner("Loading Image Model (SD3)..."):
                st.session_state.image_agent = LocalImageGenerator()
            st.success("Image Model Loaded!")
            st.session_state.chart_agent = DataVisualizer(); st.success("Data Visualizer Ready!")
        except Exception as e: st.error(f"A critical error occurred during model loading: {e}", icon="🚨")

if st.session_state.agent:
    st.header("1. Define Your Research Document")
    doc_title_input = st.text_input("Document Title:", "The Future of Artificial Intelligence")
    user_request = st.text_area("Your Request/Abstract:", "Write a comprehensive paper on the future of AI, including its potential benefits, the ethical challenges, and a speculative look at AGI. Include at least one chart and one image.", height=150)
    if st.button("✍️ Generate Initial Draft"):
        st.session_state.doc_title = doc_title_input
        with st.spinner("The agent is planning and drafting... This may take several minutes."):
            st.session_state.outline, st.session_state.document_sections, st.session_state.metrics = st.session_state.agent.run_initial_draft(user_request, st.session_state.doc_title)
        st.rerun()
else: st.warning("Please load models using the sidebar.")

if st.session_state.metrics:
    st.header("📊 Performance Metrics (Initial Draft)")
    m = st.session_state.metrics; c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total Time", f"{m['total_time']:.2f} s"); c2.metric("Outline Time", f"{m['outline_time']:.2f} s")
    c3.metric("CPU Usage", f"{psutil.cpu_percent()}%"); c4.metric("Memory Usage", f"{psutil.virtual_memory().percent}%")
    st.write(f"**Sections Generated:** {m['total_sections']}")
    with st.expander("View section times"):
        for i, t in enumerate(m['section_times']): st.write(f"- Section {i+1}: {t:.2f} seconds")

if st.session_state.document_sections:
    st.header("2. Review and Refine Draft")
    for i, section in enumerate(st.session_state.document_sections):
        with st.container(border=True):
            st.subheader(f"Section: {section['topic']}"); st.markdown(section['content'])
            feedback = st.text_area("Your feedback for this section:", key=f"feedback_{i}")
            if st.button("🔄 Regenerate Section", key=f"regen_{i}"):
                with st.spinner("The agent is revising this section..."):
                    revised_content = st.session_state.agent.regenerate_section(st.session_state.doc_title, st.session_state.outline, section['topic'], section['content'], feedback)
                    st.session_state.document_sections[i]['content'] = revised_content
                    st.rerun()

if st.session_state.document_sections:
    st.header("3. Finalize and Build Document")
    if st.button("✅ Finalize and Build .docx File", type="primary"):
        with st.spinner("Generating visuals and building document..."):
            visual_mapping, prompts = {}, []
            st.write("Parsing document for figure blocks...")
            for section in st.session_state.document_sections:
                figures = re.findall(r'(<FIGURE>.*?</FIGURE>)', section['content'], flags=re.DOTALL)
                for block in figures:
                    visual_match = re.search(r'(\[IMAGE\|(\d+)x(\d+):\s*(.*?)\]|\[CHART\|(bar|pie):\s*({.*})\])', block, re.DOTALL)
                    if visual_match:
                        full_placeholder, img_dims, img_w, img_h, img_prompt, chart_type, chart_json = visual_match.groups()
                        prompts.append({'type': 'image' if img_prompt else 'chart', 'match': full_placeholder, 'prompt': img_prompt, 'w': int(img_w) if img_w else 1024, 'h': int(img_h) if img_h else 1024, 'chart_type': chart_type, 'json': chart_json})
            st.write(f"Found {len(prompts)} visuals to generate.")
            if prompts:
                for i, v in enumerate(prompts):
                    st.write(f"- ({i+1}/{len(prompts)}) Generating {v['type']}: {v['match'][:70]}...")
                    fname = f"{st.session_state.doc_title.replace(' ', '_').lower()}_visual_{i+1}"
                    if v['type'] == 'image':
                        path = st.session_state.image_agent.generate(v['prompt'], fname, v['w'], v['h'])
                        visual_mapping[v['match']] = path
                    elif v['type'] == 'chart':
                        path = st.session_state.chart_agent.generate_chart(v['chart_type'], v['json'], fname)
                        visual_mapping[v['match']] = path
            st.write("Assembling the final .docx file...")
            filepath = save_document(st.session_state.doc_title, st.session_state.doc_title, st.session_state.document_sections, visual_mapping)
            st.success(f"Document built successfully! ✨"); st.markdown(f"**Download from:** `{filepath}`")