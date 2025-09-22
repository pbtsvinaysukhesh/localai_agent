import os
from datetime import datetime
from docx import Document
from langchain.tools import BaseTool
from pydantic import BaseModel, Field
from typing import Type, ClassVar

from config import OUTPUT_DIR

class DocumentInput(BaseModel):
    """Input model for the DocumentWriterTool."""
    filename: str = Field(description="The name of the file to create, without the .docx extension.")
    content: str = Field(description="The full content to be written into the document.")
    title: str = Field(description="The main title to be placed at the top of the document.")

class DocumentWriterTool(BaseTool):
    # --- FIX APPLIED HERE ---
    # We added the ': str' type annotation to both name and description.
    name: str = "DocumentWriter"
    description: str = "Use this tool to write a given text content into a .docx document. This should be the final step after all content has been generated and finalized."
    # --------------------------

    args_schema: Type[BaseModel] = DocumentInput

    def _run(self, filename: str, content: str, title: str):
        """Use the tool."""
        if not os.path.exists(OUTPUT_DIR):
            os.makedirs(OUTPUT_DIR)

        doc = Document()
        doc.add_heading(title, level=1)

        # Split content into paragraphs for better formatting
        for paragraph in content.split('\n'):
            if paragraph.strip():  # Avoid adding empty paragraphs
                doc.add_paragraph(paragraph)

        sanitized_filename = filename.replace(" ", "_").lower()
        filepath = os.path.join(OUTPUT_DIR, f"{sanitized_filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

        try:
            doc.save(filepath)
            return f"Document successfully saved to {filepath}."
        except Exception as e:
            return f"Error saving document: {e}"

    def _arun(self, filename: str, content: str, title: str):
        """Asynchronous version of the tool is not supported."""
        raise NotImplementedError("DocumentWriterTool does not support async")