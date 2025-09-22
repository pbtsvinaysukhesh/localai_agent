import os
import psutil
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches
from ctransformers import AutoModelForCausalLM

from config import (
    LOCAL_LLM_MODEL_PATH,
    LOCAL_LLM_MODEL_TYPE,
    LOCAL_LLM_THREADS,
    LOCAL_LLM_GPU_LAYERS,
    OUTPUT_DIR
)

class OfflineAgent:
    def __init__(self):
        print("Initializing Offline Agent...")
        self.llm = self._initialize_llm()
        self._ensure_output_directory()
        print("Offline Agent initialized.")

    def _ensure_output_directory(self):
        """Ensures the output directory exists."""
        if not os.path.exists(OUTPUT_DIR):
            os.makedirs(OUTPUT_DIR)
            print(f"Created output directory: {OUTPUT_DIR}")

    def _initialize_llm(self):
        """
        Initializes the local LLM using ctransformers.
        Checks for model file existence.
        """
        if not os.path.exists(LOCAL_LLM_MODEL_PATH):
            raise FileNotFoundError(
                f"LLM model not found at {LOCAL_LLM_MODEL_PATH}. "
                "Please download a GGUF model (e.g., from TheBloke on Hugging Face) "
                "and place it in the 'models' directory."
            )
        print(f"Loading local LLM: {LOCAL_LLM_MODEL_PATH}...")
        try:
            llm = AutoModelForCausalLM.from_pretrained(
                model_path_or_repo_id=LOCAL_LLM_MODEL_PATH,
                model_type=LOCAL_LLM_MODEL_TYPE,
                gpu_layers=LOCAL_LLM_GPU_LAYERS,
                threads=LOCAL_LLM_THREADS,
                # context_length=2048 # Adjust as needed for larger context
            )
            print("LLM loaded successfully.")
            return llm
        except Exception as e:
            print(f"Error loading LLM: {e}")
            print("Ensure ctransformers is installed correctly and your model path/type are correct.")
            print("If using GPU, ensure ctransformers was built with CUDA support.")
            raise

    def _monitor_hardware(self, task_name=""):
        """Monitors CPU, Memory, and (if applicable) GPU usage."""
        cpu_percent = psutil.cpu_percent(interval=None) # Non-blocking
        memory_info = psutil.virtual_memory()
        print(f"[{task_name}] Hardware Usage: CPU: {cpu_percent}% | Memory: {memory_info.percent}% ({memory_info.used / (1024**3):.2f} GB used)")

        # You can add GPU monitoring here using pynvml if you have NVIDIA GPUs
        # Example (requires `pip install pynvml`):
        # try:
        #     from pynvml import nvmlInit, nvmlDeviceGetHandleByIndex, nvmlDeviceGetUtilizationRates, nvmlShutdown
        #     nvmlInit()
        #     handle = nvmlDeviceGetHandleByIndex(0) # Assuming device 0
        #     utilization = nvmlDeviceGetUtilizationRates(handle)
        #     print(f"GPU Usage: {utilization.gpu}% | VRAM: {utilization.memory}%")
        #     nvmlShutdown()
        # except Exception:
        #     pass # No NVIDIA GPU or pynvml not installed/initialized

    def generate_text(self, prompt: str, max_new_tokens: int = 500, temperature: float = 0.7) -> str:
        """
        Generates text using the loaded local LLM.
        """
        self._monitor_hardware("Text Generation Start")
        print(f"\n--- Generating text with prompt ---\n'{prompt[:100]}...'")
        full_prompt = f"### System:\nYou are an AI assistant that generates high-quality content.\n\n### User:\n{prompt}\n\n### Assistant:"
        try:
            # The 'generate' method often yields tokens, so we join them.
            # Or you can call it directly for full generation if your ctransformers version allows.
            # Example using stream=True for token-by-token processing:
            # generated_tokens = []
            # for token in self.llm(full_prompt, max_new_tokens=max_new_tokens, temperature=temperature, stream=True):
            #     generated_tokens.append(token)
            # return "".join(generated_tokens)

            # More straightforward call for full generation:
            output = self.llm(full_prompt, max_new_tokens=max_new_tokens, temperature=temperature, stop=["### User:", "### System:"], stream=False)
            print(f"\n--- Text Generation Complete ---\n")
            self._monitor_hardware("Text Generation End")
            return output.strip()
        except Exception as e:
            print(f"Error during text generation: {e}")
            return "Error generating text."

    def create_document(self, filename: str, title: str, content_sections: list[str]):
        """
        Creates a .docx document with a title and multiple content sections.
        """
        self._monitor_hardware("Document Creation Start")
        doc = Document()
        doc.add_heading(title, level=1)

        for section_content in content_sections:
            doc.add_paragraph(section_content)
            doc.add_paragraph() # Add a blank line between sections

        filepath = os.path.join(OUTPUT_DIR, f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        try:
            doc.save(filepath)
            print(f"Document '{filepath}' created successfully.")
            self._monitor_hardware("Document Creation End")
            return filepath
        except Exception as e:
            print(f"Error saving document: {e}")
            return None

    def _chain_of_thought_process(self, user_request: str) -> list[str]:
        """
        Simulates a Chain of Thought process using the LLM to break down a request.
        Returns a list of generated content sections.
        """
        print(f"\n--- Starting Chain of Thought for request: '{user_request[:50]}...' ---")
        sections = []

        # Step 1: Outline Generation
        outline_prompt = f"Based on the following request, generate a detailed outline (headings and subheadings) for a document. Focus on logical flow and comprehensiveness. User Request: '{user_request}'"
        outline_text = self.generate_text(outline_prompt, max_new_tokens=300)
        print(f"\nGenerated Outline:\n{outline_text}\n")
        sections.append(f"## Document Outline\n{outline_text}")

        # Basic parsing of the outline for subsequent content generation.
        # This is a simple approach; more robust parsing might be needed for complex outlines.
        # For simplicity, let's assume the LLM provides coherent paragraphs for each section.
        # You'd typically parse headings and then generate content for each.

        # For this example, let's just ask the LLM to write directly based on the request,
        # and we'll refine the CoT later to parse the outline properly.
        # This is a placeholder for the iterative content generation based on the outline.

        print("\nNow generating content based on the outline/request...")

        # Step 2: Introduction
        intro_prompt = f"Write a comprehensive introduction for a document requested by the user. The main topic is: '{user_request}'."
        intro_content = self.generate_text(intro_prompt, max_new_tokens=400)
        sections.append(f"## Introduction\n{intro_content}")

        # Step 3: Main Body (Simplified - in a real CoT, this would iterate through outline points)
        body_prompt = f"Expand on the main points related to '{user_request}' in detail for a document. Provide several paragraphs covering different aspects. Ensure good structure and clear explanations."
        body_content = self.generate_text(body_prompt, max_new_tokens=800)
        sections.append(f"## Main Content\n{body_content}")

        # Step 4: Conclusion
        conclusion_prompt = f"Write a strong conclusion for a document discussing '{user_request}'. Summarize key findings and provide a final thought."
        conclusion_content = self.generate_text(conclusion_prompt, max_new_tokens=300)
        sections.append(f"## Conclusion\n{conclusion_content}")

        print("\n--- Chain of Thought Complete ---")
        return sections


    def run_agent(self, user_request: str, document_title: str):
        """
        Main method to run the agent's content creation process.
        """
        print(f"\nAgent received request: '{user_request}'")
        self._monitor_hardware("Agent Start")

        # Execute Chain of Thought to get content sections
        content_sections = self._chain_of_thought_process(user_request)

        # Create the final document
        if content_sections:
            filepath = self.create_document(
                filename=document_title.replace(" ", "_").lower(),
                title=document_title,
                content_sections=content_sections
            )
            if filepath:
                print(f"\nAgent successfully created document: {filepath}")
            else:
                print("\nAgent failed to create the document.")
        else:
            print("\nAgent could not generate content sections.")

        self._monitor_hardware("Agent End")
        print("\nAgent finished processing.")


# Example Usage
if __name__ == "__main__":
    # 1. Ensure you've downloaded a GGUF model and placed it in the 'models' directory.
    #    Update config.py with the correct path and type.
    # 2. Make sure you've installed ctransformers, python-docx, and psutil.

    try:
        agent = OfflineAgent()

        # Define your content creation task
        request = "Write a comprehensive report about the benefits of local AI agents for personal productivity, covering aspects like data privacy, customization, and efficiency."
        title = "The Power of Local AI for Personal Productivity"

        agent.run_agent(user_request=request, document_title=title)

    except FileNotFoundError as e:
        print(f"\nERROR: {e}")
        print("Please ensure the LLM model is correctly downloaded and configured in config.py.")
    except Exception as e:
        print(f"\nAn unexpected error occurred: {e}")
        print("Please check your ctransformers installation and hardware compatibility.")